#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SEE ALL AVKN - Quote and Invoice Management System
Application de gestion des devis et factures pour SEE ALL AVKN

Author: Expert Developer Assistant
Date: November 2025
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
import sqlite3
import datetime
import os
import re
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any
import uuid
from xml.sax.saxutils import escape

# PDF generation imports
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: reportlab not installed. PDF export will not be available.")

# Word document generation imports
try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Word export will not be available.")

# Import configuration
try:
    from config import COMPANY_CONFIG, BUSINESS_CONFIG, UI_CONFIG
    CONFIG_AVAILABLE = True
except ImportError:
    CONFIG_AVAILABLE = False
    # Default configuration
    COMPANY_CONFIG = {
        'name': 'SEE ALL AVKN',
        'address': '38 rue Dunois\n75013 PARIS',
        'email': 'michael@seeall.fr',
        'siren': '951 474 709',
        'siret': '95147470900015',
        'tva': 'FR95951474709',
        'legal_form': 'SAS',
        'capital': '1 000,00 €',
        'rcs': '951 474 709 R.C.S. Paris',
        'bank_bic': 'CMCIFRPP',
        'bank_iban': 'FR76 3006 6109 4100 0210 0820 254'
    }
    BUSINESS_CONFIG = {
        'default_vat_rate': 0.20,
        'quote_prefix': 'SA',
        'invoice_prefix': 'FA',
        'date_format': '%d/%m/%Y',
        'payment_terms': 'VIREMENT'
    }
    UI_CONFIG = {
        'window_title': 'SEE ALL AVKN - Gestion Devis & Factures',
        'window_size': '1200x800',
        'theme': 'clam'
    }

@dataclass
class Client:
    """Client data class for storing customer information"""
    id: Optional[int] = None
    name: str = ""
    siret: str = ""
    address: str = ""
    email: str = ""
    phone: str = ""
    created_at: Optional[datetime.datetime] = None

@dataclass
class SiteItem:
    """Site item data class for storing site-specific work with detailed address"""
    site_number: str = ""
    # Detailed address fields
    address: str = ""  # Street address
    postal_code: str = ""  # Code postal
    city: str = ""  # Ville
    latitude: str = ""  # Latitude (stored as string for precision)
    longitude: str = ""  # Longitude (stored as string for precision)
    # Work details
    description: str = ""
    price_ht: float = 0.0
    
    @property
    def total_ht(self) -> float:
        return self.price_ht
    
    @property
    def full_address(self) -> str:
        """Get formatted full address for display"""
        parts = []
        if self.address.strip():
            parts.append(self.address.strip())
        if self.postal_code.strip() or self.city.strip():
            city_part = f"{self.postal_code.strip()} {self.city.strip()}".strip()
            if city_part:
                parts.append(city_part)
        return ", ".join(parts) if parts else ""
    
    @property
    def coordinates(self) -> str:
        """Get formatted coordinates for display"""
        if self.latitude.strip() and self.longitude.strip():
            return f"Lat: {self.latitude}, Lng: {self.longitude}"
        return ""

@dataclass
class QuoteItem:
    """Quote/Invoice item data class - DEPRECATED, replaced by SiteItem"""
    description: str = ""
    price_ht: float = 0.0
    quantity: int = 1
    
    @property
    def total_ht(self) -> float:
        return self.price_ht * self.quantity

@dataclass
class Quote:
    """Quote data class"""
    id: Optional[int] = None
    number: str = ""
    client_id: int = 0
    client: Optional[Client] = None
    typology: str = ""  # Typologie: Pylonne, TT, CDE, Eglise, Autre
    sites: List[SiteItem] = field(default_factory=list)  # Multiple sites with their own descriptions and costs
    items: List[QuoteItem] = field(default_factory=list)  # Keep for backward compatibility
    created_at: Optional[datetime.datetime] = None
    quote_date: Optional[datetime.date] = None  # Editable quote date (different from created_at)
    intervention_date: Optional[datetime.date] = None  # Date of intervention for invoices
    is_invoice: bool = False
    invoice_number: Optional[str] = None
    order_number: str = ""  # Bon de commande number for invoices
    site_number: str = ""  # DEPRECATED - kept for migration compatibility
    site_address: str = ""  # DEPRECATED - kept for migration compatibility
    is_invoiced: bool = False  # True if this quote has been converted to an invoice
    linked_invoice_id: Optional[int] = None  # ID of the corresponding invoice
    
    @property
    def total_ht(self) -> float:
        # Calculate total from sites (new way) or items (old way)
        sites_total = sum(site.total_ht for site in self.sites)
        items_total = sum(item.total_ht for item in self.items)
        return sites_total + items_total
    
    @property
    def total_tva(self) -> float:
        vat_rate = BUSINESS_CONFIG.get('default_vat_rate', 0.20)
        return self.total_ht * vat_rate
    
    @property
    def total_ttc(self) -> float:
        return self.total_ht + self.total_tva
    
    @property
    def site_numbers_list(self) -> List[str]:
        """Get list of all site numbers for this quote"""
        site_numbers = [site.site_number for site in self.sites if site.site_number.strip()]
        # Include legacy site_number if exists and not already in list
        if self.site_number.strip() and self.site_number not in site_numbers:
            site_numbers.append(self.site_number)
        return site_numbers
    
    @property
    def site_numbers_display(self) -> str:
        """Get formatted string of all site numbers for display"""
        numbers = self.site_numbers_list
        if not numbers:
            return "Aucun site"
        elif len(numbers) == 1:
            return numbers[0]
        else:
            return f"{numbers[0]} (+{len(numbers)-1})"

class DatabaseManager:
    """Database manager class for handling SQLite operations"""
    
    def __init__(self, db_path: str = None):
        if db_path is None:
            db_path = 'seeall_database.db'
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """Initialize database with required tables"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Create clients table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS clients (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT NOT NULL,
                    siret TEXT,
                    address TEXT,
                    email TEXT,
                    phone TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # Create quotes table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS quotes (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    number TEXT UNIQUE NOT NULL,
                    client_id INTEGER,
                    typology TEXT,
                    quote_date DATE,
                    intervention_date DATE,
                    is_invoice BOOLEAN DEFAULT FALSE,
                    invoice_number TEXT,
                    order_number TEXT,
                    site_number TEXT,
                    site_address TEXT,
                    is_invoiced BOOLEAN DEFAULT FALSE,
                    linked_invoice_id INTEGER,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (client_id) REFERENCES clients (id),
                    FOREIGN KEY (linked_invoice_id) REFERENCES quotes (id)
                )
            ''')
            
            # Create sites table for multiple sites per quote
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS quote_sites (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    quote_id INTEGER,
                    site_number TEXT NOT NULL,
                    address TEXT,
                    postal_code TEXT,
                    city TEXT,
                    latitude TEXT,
                    longitude TEXT,
                    description TEXT NOT NULL,
                    price_ht REAL NOT NULL,
                    FOREIGN KEY (quote_id) REFERENCES quotes (id)
                )
            ''')
            
            # Create quote_items table (keep for backward compatibility)
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS quote_items (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    quote_id INTEGER,
                    description TEXT NOT NULL,
                    price_ht REAL NOT NULL,
                    quantity INTEGER DEFAULT 1,
                    FOREIGN KEY (quote_id) REFERENCES quotes (id)
                )
            ''')
            
            # Create counters table for auto-increment numbers
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS counters (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    client_name TEXT,
                    month_year TEXT,
                    counter INTEGER DEFAULT 0,
                    UNIQUE(client_name, month_year)
                )
            ''')
            
            conn.commit()
            
            # Migration: Add quote_date column if it doesn't exist (for existing databases)
            try:
                cursor.execute('ALTER TABLE quotes ADD COLUMN quote_date DATE')
                conn.commit()
                print("Migration: Added quote_date column to quotes table")
            except sqlite3.OperationalError:
                # Column already exists, migration not needed
                pass
            
            # Migration: Add intervention_date column if it doesn't exist (for existing databases)
            try:
                cursor.execute('ALTER TABLE quotes ADD COLUMN intervention_date DATE')
                conn.commit()
                print("Migration: Added intervention_date column to quotes table")
            except sqlite3.OperationalError:
                # Column already exists, migration not needed
                pass
            
            # Migration: Add typology column if it doesn't exist (for existing databases)
            try:
                cursor.execute('ALTER TABLE quotes ADD COLUMN typology TEXT')
                conn.commit()
                print("Migration: Added typology column to quotes table")
            except sqlite3.OperationalError:
                # Column already exists, migration not needed
                pass
            
            # Migration: Add detailed address columns to quote_sites if they don't exist
            new_site_columns = [
                ('address', 'TEXT'),
                ('postal_code', 'TEXT'),
                ('city', 'TEXT'),
                ('latitude', 'TEXT'),
                ('longitude', 'TEXT')
            ]
            
            for column_name, column_type in new_site_columns:
                try:
                    cursor.execute(f'ALTER TABLE quote_sites ADD COLUMN {column_name} {column_type}')
                    conn.commit()
                    print(f"Migration: Added {column_name} column to quote_sites table")
                except sqlite3.OperationalError:
                    # Column already exists, migration not needed
                    pass
            
            # Migration: Add site columns if they don't exist (for existing databases)
            try:
                cursor.execute("ALTER TABLE quotes ADD COLUMN site_number TEXT DEFAULT ''")
                cursor.execute("ALTER TABLE quotes ADD COLUMN site_address TEXT DEFAULT ''")
                cursor.execute("ALTER TABLE quotes ADD COLUMN is_invoiced BOOLEAN DEFAULT FALSE")
                cursor.execute("ALTER TABLE quotes ADD COLUMN linked_invoice_id INTEGER")
                conn.commit()
            except sqlite3.OperationalError:
                # Columns already exist
                pass
            
            # Create new sites table if it doesn't exist
            try:
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS quote_sites (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        quote_id INTEGER,
                        site_number TEXT NOT NULL,
                        site_address TEXT,
                        description TEXT NOT NULL,
                        price_ht REAL NOT NULL,
                        FOREIGN KEY (quote_id) REFERENCES quotes (id)
                    )
                ''')
                conn.commit()
            except sqlite3.OperationalError:
                # Table already exists
                pass
    
    def add_client(self, client: Client) -> int:
        """Add a new client to the database"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO clients (name, siret, address, email, phone)
                VALUES (?, ?, ?, ?, ?)
            ''', (client.name, client.siret, client.address, client.email, client.phone))
            return cursor.lastrowid
    
    def get_clients(self) -> List[Client]:
        """Get all clients from database"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM clients ORDER BY name')
            rows = cursor.fetchall()
            
            clients = []
            for row in rows:
                client = Client(
                    id=row[0],
                    name=row[1],
                    siret=row[2],
                    address=row[3],
                    email=row[4],
                    phone=row[5],
                    created_at=datetime.datetime.fromisoformat(row[6]) if row[6] else None
                )
                clients.append(client)
            return clients
    
    def get_client_by_id(self, client_id: int) -> Optional[Client]:
        """Get client by ID"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM clients WHERE id = ?', (client_id,))
            row = cursor.fetchone()
            
            if row:
                return Client(
                    id=row[0],
                    name=row[1],
                    siret=row[2],
                    address=row[3],
                    email=row[4],
                    phone=row[5],
                    created_at=datetime.datetime.fromisoformat(row[6]) if row[6] else None
                )
            return None
    
    def generate_quote_number(self, client_name: str, is_invoice: bool = False) -> str:
        """Generate automatic quote number: <PREFIX>.<client_name>.MMYYYY001"""
        # Get prefix from configuration
        prefix = BUSINESS_CONFIG.get('invoice_prefix' if is_invoice else 'quote_prefix', 
                                   'FA' if is_invoice else 'SA')
        
        # Clean client name for use in quote number (remove spaces, special chars)
        clean_name = re.sub(r'[^A-Za-z0-9]', '', client_name.upper())[:10]
        
        # Get current month/year
        now = datetime.datetime.now()
        month_year = now.strftime("%m%Y")
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Get or create counter for this client/month combination
            cursor.execute('''
                INSERT OR IGNORE INTO counters (client_name, month_year, counter)
                VALUES (?, ?, 0)
            ''', (clean_name, month_year))
            
            # Increment counter
            cursor.execute('''
                UPDATE counters SET counter = counter + 1
                WHERE client_name = ? AND month_year = ?
            ''', (clean_name, month_year))
            
            # Get current counter value
            cursor.execute('''
                SELECT counter FROM counters
                WHERE client_name = ? AND month_year = ?
            ''', (clean_name, month_year))
            
            counter = cursor.fetchone()[0]
            conn.commit()
        
        # Format: <PREFIX>.<CLIENT>.MMYYYY001
        return f"{prefix}.{clean_name}.{month_year}{counter:03d}"
    
    def save_quote(self, quote: Quote) -> int:
        """Save quote to database (create new or update existing)"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Convert dates to string format for SQLite
            quote_date_str = quote.quote_date.isoformat() if quote.quote_date else None
            intervention_date_str = quote.intervention_date.isoformat() if quote.intervention_date else None
            
            if quote.id is None:
                # Create new quote
                cursor.execute('''
                    INSERT INTO quotes (number, client_id, typology, quote_date, intervention_date, is_invoice, 
                                      invoice_number, order_number, site_number, site_address, 
                                      is_invoiced, linked_invoice_id)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (quote.number, quote.client_id, quote.typology, quote_date_str, intervention_date_str, 
                      quote.is_invoice, quote.invoice_number, quote.order_number, 
                      quote.site_number, quote.site_address, quote.is_invoiced, quote.linked_invoice_id))
                
                quote_id = cursor.lastrowid
                quote.id = quote_id  # Set the ID for the quote object
            else:
                # Update existing quote
                cursor.execute('''
                    UPDATE quotes SET client_id=?, typology=?, quote_date=?, intervention_date=?, 
                                    invoice_number=?, order_number=?, site_number=?, site_address=?, 
                                    is_invoiced=?, linked_invoice_id=?
                    WHERE id=?
                ''', (quote.client_id, quote.typology, quote_date_str, intervention_date_str,
                      quote.invoice_number, quote.order_number, quote.site_number, quote.site_address,
                      quote.is_invoiced, quote.linked_invoice_id, quote.id))
                
                quote_id = quote.id
                
                # Delete existing sites and items before inserting updated ones
                cursor.execute('DELETE FROM quote_sites WHERE quote_id = ?', (quote_id,))
                cursor.execute('DELETE FROM quote_items WHERE quote_id = ?', (quote_id,))
            
            # Insert/Re-insert quote sites (new multi-site system)
            for site in quote.sites:
                cursor.execute('''
                    INSERT INTO quote_sites (quote_id, site_number, address, postal_code, city, 
                                           latitude, longitude, description, price_ht)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (quote_id, site.site_number, site.address, site.postal_code, site.city,
                      site.latitude, site.longitude, site.description, site.price_ht))
            
            # Insert/Re-insert quote items (for backward compatibility)
            for item in quote.items:
                cursor.execute('''
                    INSERT INTO quote_items (quote_id, description, price_ht, quantity)
                    VALUES (?, ?, ?, ?)
                ''', (quote_id, item.description, item.price_ht, item.quantity))
            
            conn.commit()
            return quote_id
    
    def get_quotes(self, is_invoice: bool = False) -> List[Quote]:
        """Get all quotes or invoices from database"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT q.id, q.number, q.client_id, q.typology, q.quote_date, q.intervention_date, 
                       q.is_invoice, q.invoice_number, q.order_number, q.site_number, 
                       q.site_address, q.is_invoiced, q.linked_invoice_id, q.created_at,
                       c.name as client_name
                FROM quotes q
                LEFT JOIN clients c ON q.client_id = c.id
                WHERE q.is_invoice = ?
                ORDER BY q.created_at DESC
            ''', (is_invoice,))
            
            quotes = []
            for row in cursor.fetchall():
                # Parse dates from ISO format strings
                quote_date = None
                if row[4]:  # quote_date (index shifted due to typology)
                    try:
                        quote_date = datetime.date.fromisoformat(row[4])
                    except (ValueError, TypeError):
                        quote_date = None
                
                intervention_date = None
                if row[5]:  # intervention_date (index shifted due to typology)
                    try:
                        intervention_date = datetime.date.fromisoformat(row[5])
                    except (ValueError, TypeError):
                        intervention_date = None
                
                created_at = None
                if row[13]:  # created_at (index shifted due to typology)
                    try:
                        created_at = datetime.datetime.fromisoformat(row[13])
                    except (ValueError, TypeError):
                        created_at = None
                
                quote = Quote(
                    id=row[0],
                    number=row[1],
                    client_id=row[2],
                    typology=row[3] or "",  # New typology field
                    quote_date=quote_date,
                    intervention_date=intervention_date,
                    is_invoice=bool(row[6]),
                    invoice_number=row[7],
                    order_number=row[8],
                    site_number=row[9] or "",
                    site_address=row[10] or "",
                    is_invoiced=bool(row[11]) if row[11] is not None else False,
                    linked_invoice_id=row[12] if row[12] else None,
                    created_at=created_at
                )
                
                # Get client info
                if quote.client_id:
                    quote.client = self.get_client_by_id(quote.client_id)
                
                # Get quote sites (new multi-site system)
                cursor.execute('''
                    SELECT site_number, address, postal_code, city, latitude, longitude, description, price_ht
                    FROM quote_sites
                    WHERE quote_id = ?
                ''', (quote.id,))
                
                for site_row in cursor.fetchall():
                    site = SiteItem(
                        site_number=site_row[0] or "",
                        address=site_row[1] or "",
                        postal_code=site_row[2] or "",
                        city=site_row[3] or "",
                        latitude=site_row[4] or "",
                        longitude=site_row[5] or "",
                        description=site_row[6] or "",
                        price_ht=site_row[7] or 0.0
                    )
                    quote.sites.append(site)
                
                # Get quote items (for backward compatibility)
                cursor.execute('''
                    SELECT description, price_ht, quantity
                    FROM quote_items
                    WHERE quote_id = ?
                ''', (quote.id,))
                
                for item_row in cursor.fetchall():
                    item = QuoteItem(
                        description=item_row[0],
                        price_ht=item_row[1],
                        quantity=item_row[2]
                    )
                    quote.items.append(item)
                
                quotes.append(quote)
            
            return quotes
    
    def delete_client(self, client_id: int) -> bool:
        """Delete a client from database
        
        Args:
            client_id: ID of the client to delete
            
        Returns:
            bool: True if deletion successful, False otherwise
            
        Raises:
            ValueError: If client has associated quotes/invoices
        """
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Check if client exists
            cursor.execute('SELECT name FROM clients WHERE id = ?', (client_id,))
            client_data = cursor.fetchone()
            if not client_data:
                raise ValueError("Client not found")
            
            # Check for associated quotes/invoices - BLOCK deletion if any exist
            cursor.execute('SELECT COUNT(*) FROM quotes WHERE client_id = ?', (client_id,))
            quote_count = cursor.fetchone()[0]
            if quote_count > 0:
                raise ValueError(f"Impossible de supprimer le client '{client_data[0]}' : {quote_count} devis/facture(s) associé(s). Veuillez d'abord supprimer tous les devis et factures de ce client.")
            
            # Delete the client (no associated data to worry about)
            cursor.execute('DELETE FROM clients WHERE id = ?', (client_id,))
            conn.commit()
            
            return True
    
    def delete_quote(self, quote_id: int) -> bool:
        """Delete a quote or invoice from database
        
        Args:
            quote_id: ID of the quote/invoice to delete
            
        Returns:
            bool: True if deletion successful, False otherwise
            
        Raises:
            ValueError: If quote not found or if trying to delete an invoiced quote
        """
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Get quote info
            cursor.execute('SELECT number, is_invoice, linked_invoice_id, is_invoiced FROM quotes WHERE id = ?', (quote_id,))
            quote_data = cursor.fetchone()
            if not quote_data:
                raise ValueError("Devis/Facture introuvable")
            
            quote_number, is_invoice, linked_invoice_id, is_invoiced = quote_data
            
            # BLOCK deletion if this is a quote that has been invoiced
            if not is_invoice and is_invoiced:
                raise ValueError(f"Impossible de supprimer le devis '{quote_number}' : il a été converti en facture. Veuillez d'abord supprimer la facture correspondante.")
            
            # If this is an invoice, update the original quote to mark it as not invoiced
            if is_invoice:
                cursor.execute('''
                    UPDATE quotes 
                    SET is_invoiced = FALSE, linked_invoice_id = NULL 
                    WHERE linked_invoice_id = ?
                ''', (quote_id,))
            
            # Delete quote items
            cursor.execute('DELETE FROM quote_items WHERE quote_id = ?', (quote_id,))
            
            # Delete quote sites
            cursor.execute('DELETE FROM quote_sites WHERE quote_id = ?', (quote_id,))
            
            # Delete the quote/invoice
            cursor.execute('DELETE FROM quotes WHERE id = ?', (quote_id,))
            conn.commit()
            
            return True
    
    def get_client_quote_count(self, client_id: int) -> int:
        """Get the number of quotes/invoices for a client"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT COUNT(*) FROM quotes WHERE client_id = ?', (client_id,))
            return cursor.fetchone()[0]

    def convert_quote_to_invoice(self, quote_id: int, order_number: str, intervention_date: datetime.date = None) -> str:
        """Convert a quote to an invoice by creating a new invoice entry and marking the original quote as invoiced"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Get original quote data
            cursor.execute('''
                SELECT id, number, client_id, typology, quote_date, intervention_date, is_invoice, 
                       invoice_number, order_number, site_number, site_address, 
                       is_invoiced, linked_invoice_id, created_at
                FROM quotes WHERE id = ?
            ''', (quote_id,))
            quote_row = cursor.fetchone()
            
            if not quote_row:
                raise ValueError("Quote not found")
            
            # Get client info for invoice number generation
            client = self.get_client_by_id(quote_row[2])  # client_id is at index 2
            
            # Generate invoice number using the same pattern but with invoice prefix
            invoice_number = self.generate_quote_number(client.name, is_invoice=True)
            
            # Convert intervention_date to string format for SQLite
            intervention_date_str = intervention_date.isoformat() if intervention_date else None
            
            # Create new invoice entry (copy from quote but as invoice)
            cursor.execute('''
                INSERT INTO quotes (number, client_id, typology, quote_date, intervention_date, is_invoice, 
                                  invoice_number, order_number, site_number, site_address, 
                                  is_invoiced, linked_invoice_id)
                VALUES (?, ?, ?, ?, ?, TRUE, ?, ?, ?, ?, FALSE, NULL)
            ''', (invoice_number, quote_row[2], quote_row[3], quote_row[4], intervention_date_str, 
                  invoice_number, order_number, quote_row[9] or "", quote_row[10] or ""))
            
            invoice_id = cursor.lastrowid
            
            # Copy quote sites to the new invoice
            cursor.execute('''
                SELECT site_number, address, postal_code, city, latitude, longitude, description, price_ht 
                FROM quote_sites WHERE quote_id = ?
            ''', (quote_id,))
            sites = cursor.fetchall()
            
            for site in sites:
                cursor.execute('''
                    INSERT INTO quote_sites (quote_id, site_number, address, postal_code, city, 
                                           latitude, longitude, description, price_ht)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (invoice_id, site[0], site[1], site[2], site[3], site[4], site[5], site[6], site[7]))
            
            # Copy quote items to the new invoice (for backward compatibility)
            cursor.execute('SELECT description, price_ht, quantity FROM quote_items WHERE quote_id = ?', (quote_id,))
            items = cursor.fetchall()
            
            for item in items:
                cursor.execute('''
                    INSERT INTO quote_items (quote_id, description, price_ht, quantity)
                    VALUES (?, ?, ?, ?)
                ''', (invoice_id, item[0], item[1], item[2]))
            
            # Mark original quote as invoiced and link to the new invoice
            cursor.execute('''
                UPDATE quotes 
                SET is_invoiced = TRUE, linked_invoice_id = ?
                WHERE id = ?
            ''', (invoice_id, quote_id))
            
            conn.commit()
            return invoice_number

class PDFGenerator:
    """PDF generator class using reportlab"""
    
    @staticmethod
    def generate_quote_pdf(quote: Quote, filepath: str):
        """Generate PDF for quote or invoice"""
        if not PDF_AVAILABLE:
            raise ImportError("reportlab not available. Cannot generate PDF documents.")
        
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()

        description_style = ParagraphStyle(
            'DescriptionStyle',
            parent=styles['Normal'],
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            spaceAfter=4,
            wordWrap='CJK'
        )

        def build_description_paragraph(lines: List[str]) -> Paragraph:
            """Build a paragraph that keeps long text inside the table cell."""
            processed = []
            for line in lines:
                if not line:
                    continue
                processed.append(escape(line).replace('\n', '<br/>'))
            if not processed:
                processed.append('')
            return Paragraph('<br/>'.join(processed), description_style)
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=16,
            textColor=colors.black,
            alignment=TA_CENTER
        )
        
        # Document title
        title = "FACTURE" if quote.is_invoice else "DEVIS"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # Determine the appropriate date to display
        if quote.is_invoice:
            # For invoices, use intervention_date if available, otherwise created_at
            if quote.intervention_date:
                display_date = quote.intervention_date.strftime(BUSINESS_CONFIG.get('date_format', '%d/%m/%Y'))
            elif quote.created_at:
                display_date = quote.created_at.strftime(BUSINESS_CONFIG.get('date_format', '%d/%m/%Y'))
            else:
                display_date = datetime.datetime.now().strftime(BUSINESS_CONFIG.get('date_format', '%d/%m/%Y'))
        else:
            # For quotes, use quote_date if available, otherwise created_at
            if quote.quote_date:
                display_date = quote.quote_date.strftime(BUSINESS_CONFIG.get('date_format', '%d/%m/%Y'))
            elif quote.created_at:
                display_date = quote.created_at.strftime(BUSINESS_CONFIG.get('date_format', '%d/%m/%Y'))
            else:
                display_date = datetime.datetime.now().strftime(BUSINESS_CONFIG.get('date_format', '%d/%m/%Y'))
        
        # Company info table
        company_data = [
            [COMPANY_CONFIG['name'], f"Date: {display_date}"],
            [COMPANY_CONFIG['address'].split('\n')[0], ""],
            [COMPANY_CONFIG['address'].split('\n')[1] if len(COMPANY_CONFIG['address'].split('\n')) > 1 else "", ""],
            [f"Email: {COMPANY_CONFIG['email']}", ""],
            [f"TVA: {COMPANY_CONFIG['tva']}", ""]
        ]
        
        company_table = Table(company_data, colWidths=[10*cm, 8*cm])
        company_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
        ]))
        story.append(company_table)
        story.append(Spacer(1, 20))
        
        # Document number and client info
        doc_number = quote.invoice_number if quote.is_invoice else quote.number
        info_data = [
            [f"Numéro: {doc_number}", ""],
            ["", ""],
            ["Client:", ""],
            [quote.client.name if quote.client else "", ""],
            [quote.client.address if quote.client else "", ""],
            [f"SIRET: {quote.client.siret}" if quote.client and quote.client.siret else "", ""],
            ["", ""]
        ]
        
        # Add order number for invoices
        if quote.is_invoice and quote.order_number:
            info_data.insert(1, [f"Bon de commande: {quote.order_number}", ""])
        
        # Add sites information - show all sites or legacy site
        if quote.sites:
            # New multi-site system
            info_data.append(["Sites d'intervention:", ""])
            for i, site in enumerate(quote.sites[:3]):  # Show first 3 sites, add "..." if more
                site_info = f"Site {site.site_number}"
                if site.full_address:
                    site_info += f" - {site.full_address}"
                info_data.append([site_info, ""])
            
            if len(quote.sites) > 3:
                info_data.append([f"... et {len(quote.sites) - 3} autre(s) site(s)", ""])
        elif quote.site_number or quote.site_address:
            # Legacy single site system
            info_data.extend([
                ["Site d'intervention:", ""],
                [f"Numéro de site: {quote.site_number}" if quote.site_number else "", ""],
                [quote.site_address if quote.site_address else "", ""]
            ])
        
        info_table = Table(info_data, colWidths=[10*cm, 8*cm])
        info_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 2), (0, 2), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 30))
        
        # Items table - Include both sites and legacy items
        items_data = [["Description", "Prix HT"]]

        # Add sites (new system) with detailed descriptions
        for site in quote.sites:
            site_header = "Site"
            if site.site_number:
                site_header += f" {site.site_number}"
            if site.full_address:
                site_header += f" - {site.full_address}"
            if site.coordinates:
                site_header += f" ({site.coordinates})"

            description_lines = [site_header]
            if site.description:
                description_lines.append(site.description)

            items_data.append([
                build_description_paragraph(description_lines),
                f"{site.total_ht:.2f} €"
            ])

        # Add legacy items (for backward compatibility)
        for item in quote.items:
            items_data.append([
                build_description_paragraph([item.description]),
                f"{item.total_ht:.2f} €"
            ])
        items_table = Table(items_data, colWidths=[14*cm, 4*cm])
        items_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align content to top for multi-line descriptions
        ]))
        story.append(items_table)
        story.append(Spacer(1, 20))
        
        # Totals table
        totals_data = [
            ["Montant total HT:", f"{quote.total_ht:.2f} €"],
            ["Montant total TVA:", f"{quote.total_tva:.2f}€"],
            ["Montant total TTC:", f"{quote.total_ttc:.2f} €"],
            ["NET A PAYER:", f"{quote.total_ttc:.2f} €"]
        ]
        
        totals_table = Table(totals_data, colWidths=[14*cm, 4*cm])
        totals_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
        ]))
        story.append(totals_table)
        story.append(Spacer(1, 30))
        
        # Payment info for invoices
        if quote.is_invoice:
            payment_method = BUSINESS_CONFIG.get('payment_terms', 'VIREMENT')
            payment_info = f"""
            Mode de règlement par {payment_method}
            BIC : {COMPANY_CONFIG['bank_bic']} IBAN : {COMPANY_CONFIG['bank_iban']}
            
            (En cas de retard de paiement, une pénalité égale à trois fois le taux d'intérêt légal sera exigible et 
            une indemnité pour frais de recouvrement de 40€ sera appliquée article L.441-6). Nos conditions de 
            ventes ne prévoient pas d'escompte en cas de paiement anticipé.)
            """
            story.append(Paragraph(payment_info, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Legal info
        legal_info = f"""
        La Société dénommée {COMPANY_CONFIG['name']}, {COMPANY_CONFIG['legal_form']}, au capital social de {COMPANY_CONFIG['capital']}, 
        inscrit sous le numéro de Siren {COMPANY_CONFIG['siren']}/ Siret n°{COMPANY_CONFIG['siret']} au {COMPANY_CONFIG['rcs']}
        """
        story.append(Paragraph(legal_info, styles['Normal']))
        
        # Build PDF
        doc.build(story)

class WordGenerator:
    """Word document generator class"""
    
    @staticmethod
    def generate_quote_docx(quote: Quote, filepath: str):
        """Generate Word document for quote or invoice"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Cannot generate Word documents.")
        
        doc = Document()
        
        # Document title
        title = "FACTURE" if quote.is_invoice else "DEVIS"
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add space
        doc.add_paragraph()
        
        # Company info
        company_para = doc.add_paragraph()
        company_para.add_run(COMPANY_CONFIG['name']).bold = True
        company_para.add_run(f"\n{COMPANY_CONFIG['address']}\nEmail: {COMPANY_CONFIG['email']}\nTVA: {COMPANY_CONFIG['tva']}")
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_format = BUSINESS_CONFIG.get('date_format', '%d/%m/%Y')
        
        # Determine the appropriate date to display
        if quote.is_invoice:
            # For invoices, use intervention_date if available, otherwise created_at
            if quote.intervention_date:
                date_str = quote.intervention_date.strftime(date_format)
            elif quote.created_at:
                date_str = quote.created_at.strftime(date_format)
            else:
                date_str = datetime.datetime.now().strftime(date_format)
        else:
            # For quotes, use quote_date if available, otherwise created_at
            if quote.quote_date:
                date_str = quote.quote_date.strftime(date_format)
            elif quote.created_at:
                date_str = quote.created_at.strftime(date_format)
            else:
                date_str = datetime.datetime.now().strftime(date_format)
        
        date_para.add_run(f"Date: {date_str}")
        
        doc.add_paragraph()
        
        # Document number
        doc_number = quote.invoice_number if quote.is_invoice else quote.number
        number_para = doc.add_paragraph()
        number_para.add_run(f"Numéro: {doc_number}").bold = True
        
        # Order number for invoices
        if quote.is_invoice and quote.order_number:
            order_para = doc.add_paragraph()
            order_para.add_run(f"Bon de commande: {quote.order_number}").bold = True
        
        doc.add_paragraph()
        
        # Client info
        client_para = doc.add_paragraph()
        client_para.add_run("Client:").bold = True
        if quote.client:
            client_para.add_run(f"\n{quote.client.name}")
            if quote.client.address:
                client_para.add_run(f"\n{quote.client.address}")
            if quote.client.siret:
                client_para.add_run(f"\nSIRET: {quote.client.siret}")
        
        # Site information - show all sites or legacy site
        if quote.sites:
            # New multi-site system
            site_para = doc.add_paragraph()
            site_para.add_run("Sites d'intervention:").bold = True
            
            for site in quote.sites:
                site_info = f"\nSite {site.site_number}"
                if site.full_address:
                    site_info += f" - {site.full_address}"
                if site.coordinates:
                    site_info += f" ({site.coordinates})"
                site_para.add_run(site_info)
        elif quote.site_number or quote.site_address:
            # Legacy single site system
            site_para = doc.add_paragraph()
            site_para.add_run("Site d'intervention:").bold = True
            if quote.site_number:
                site_para.add_run(f"\nNuméro de site: {quote.site_number}")
            if quote.site_address:
                site_para.add_run(f"\n{quote.site_address}")
        
        doc.add_paragraph()
        
        # Items table - Include both sites and legacy items
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Description'
        header_cells[1].text = 'Prix HT'
        
        # Make header bold
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add sites (new system) with detailed descriptions
        for site in quote.sites:
            row_cells = table.add_row().cells
            
            # Build comprehensive description including site details
            site_description = f"Site {site.site_number}"
            if site.full_address:
                site_description += f" - {site.full_address}"
            if site.coordinates:
                site_description += f" ({site.coordinates})"
            if site.description:
                site_description += f"\n{site.description}"
            
            row_cells[0].text = site_description
            row_cells[1].text = f"{site.total_ht:.2f}€"
        
        # Add legacy items (for backward compatibility)
        for item in quote.items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.description
            row_cells[1].text = f"{item.total_ht:.2f}€"
        
        doc.add_paragraph()
        
        # Totals
        totals_para = doc.add_paragraph()
        totals_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        totals_para.add_run(f"Montant total HT: {quote.total_ht:.2f} €\n")
        totals_para.add_run(f"Montant total TVA: {quote.total_tva:.2f}€\n")
        totals_para.add_run(f"Montant total TTC: {quote.total_ttc:.2f} €\n")
        totals_para.add_run(f"NET A PAYER: {quote.total_ttc:.2f} €").bold = True
        
        doc.add_paragraph()
        
        # Payment info for invoices
        if quote.is_invoice:
            payment_para = doc.add_paragraph()
            payment_method = BUSINESS_CONFIG.get('payment_terms', 'VIREMENT')
            payment_para.add_run(f"Mode de règlement par {payment_method}").bold = True
            payment_para.add_run(f"\nBIC : {COMPANY_CONFIG['bank_bic']} IBAN : {COMPANY_CONFIG['bank_iban']}")
            
            payment_conditions = "(En cas de retard de paiement, une pénalité égale à trois fois le taux d'intérêt légal sera exigible et une indemnité pour frais de recouvrement de 40€ sera appliquée article L.441-6). Nos conditions de ventes ne prévoient pas d'escompte en cas de paiement anticipé.)"
            payment_para.add_run(f"\n\n{payment_conditions}")
        
        doc.add_paragraph()
        
        # Legal info
        legal_para = doc.add_paragraph()
        legal_text = f"La Société dénommée {COMPANY_CONFIG['name']}, {COMPANY_CONFIG['legal_form']}, au capital social de {COMPANY_CONFIG['capital']}, inscrit sous le numéro de Siren {COMPANY_CONFIG['siren']}/ Siret n°{COMPANY_CONFIG['siret']} au {COMPANY_CONFIG['rcs']}"
        legal_para.add_run(legal_text)
        
        # Save document
        doc.save(filepath)

class MainApplication:
    """Main application class with GUI"""
    
    def __init__(self):
        self.db = DatabaseManager()
        self.root = tk.Tk()
        
        # Use configuration for window settings
        window_title = UI_CONFIG.get('window_title', 'SEE ALL AVKN - Gestion Devis & Factures')
        window_size = UI_CONFIG.get('window_size', '1200x800')
        theme = UI_CONFIG.get('theme', 'clam')
        
        self.root.title(window_title)
        self.root.geometry(window_size)
        
        # Style configuration
        style = ttk.Style()
        try:
            style.theme_use(theme)
        except tk.TclError:
            # Fallback theme if specified theme is not available
            style.theme_use('default')
        
        self.setup_ui()
        
    def setup_ui(self):
        """Setup the user interface"""
        # Create notebook for tabs
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Clients tab
        self.clients_frame = ttk.Frame(notebook)
        notebook.add(self.clients_frame, text="Clients")
        self.setup_clients_tab()
        
        # Quotes tab
        self.quotes_frame = ttk.Frame(notebook)
        notebook.add(self.quotes_frame, text="Devis")
        self.setup_quotes_tab()
        
        # Invoices tab
        self.invoices_frame = ttk.Frame(notebook)
        notebook.add(self.invoices_frame, text="Factures")
        self.setup_invoices_tab()
    
    def setup_clients_tab(self):
        """Setup clients management tab"""
        # Title
        title_label = ttk.Label(self.clients_frame, text="Gestion des Clients", font=('Arial', 16, 'bold'))
        title_label.pack(pady=10)
        
        # Add client form
        form_frame = ttk.LabelFrame(self.clients_frame, text="Ajouter un client", padding=10)
        form_frame.pack(fill='x', padx=10, pady=5)
        
        # Form fields
        fields_frame = ttk.Frame(form_frame)
        fields_frame.pack(fill='x')
        
        # Name
        ttk.Label(fields_frame, text="Nom *:").grid(row=0, column=0, sticky='w', padx=(0, 5))
        self.client_name_var = tk.StringVar()
        ttk.Entry(fields_frame, textvariable=self.client_name_var, width=40).grid(row=0, column=1, sticky='ew', padx=(0, 10))
        
        # SIRET
        ttk.Label(fields_frame, text="SIRET:").grid(row=0, column=2, sticky='w', padx=(0, 5))
        self.client_siret_var = tk.StringVar()
        ttk.Entry(fields_frame, textvariable=self.client_siret_var, width=20).grid(row=0, column=3, sticky='ew')
        
        # Address
        ttk.Label(fields_frame, text="Adresse:").grid(row=1, column=0, sticky='nw', padx=(0, 5), pady=(5, 0))
        self.client_address_var = tk.StringVar()
        address_entry = tk.Text(fields_frame, height=3, width=40)
        address_entry.grid(row=1, column=1, columnspan=2, sticky='ew', padx=(0, 10), pady=(5, 0))
        
        # Email and Phone
        ttk.Label(fields_frame, text="Email:").grid(row=2, column=0, sticky='w', padx=(0, 5), pady=(5, 0))
        self.client_email_var = tk.StringVar()
        ttk.Entry(fields_frame, textvariable=self.client_email_var, width=40).grid(row=2, column=1, sticky='ew', padx=(0, 10), pady=(5, 0))
        
        ttk.Label(fields_frame, text="Téléphone:").grid(row=2, column=2, sticky='w', padx=(0, 5), pady=(5, 0))
        self.client_phone_var = tk.StringVar()
        ttk.Entry(fields_frame, textvariable=self.client_phone_var, width=20).grid(row=2, column=3, sticky='ew', pady=(5, 0))
        
        # Configure grid weights
        fields_frame.columnconfigure(1, weight=1)
        
        # Store address text widget reference
        self.client_address_text = address_entry
        
        # Add button
        add_button = ttk.Button(form_frame, text="Ajouter Client", command=self.add_client)
        add_button.pack(pady=(10, 0))
        
        # Clients list
        list_frame = ttk.LabelFrame(self.clients_frame, text="Clients existants", padding=10)
        list_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        # Buttons frame for client actions
        clients_buttons_frame = ttk.Frame(list_frame)
        clients_buttons_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Button(clients_buttons_frame, text="Supprimer Client", 
                  command=self.delete_selected_client).pack(side='left', padx=(0, 10))
        ttk.Button(clients_buttons_frame, text="Modifier Client", 
                  command=self.edit_selected_client).pack(side='left')
        
        # Treeview for clients list
        columns = ('ID', 'Nom', 'SIRET', 'Email', 'Téléphone')
        self.clients_tree = ttk.Treeview(list_frame, columns=columns, show='headings')
        
        for col in columns:
            self.clients_tree.heading(col, text=col)
            self.clients_tree.column(col, width=120)
        
        # Context menu for clients
        self.clients_context_menu = tk.Menu(self.root, tearoff=0)
        self.clients_context_menu.add_command(label="Modifier", command=self.edit_selected_client)
        self.clients_context_menu.add_separator()
        self.clients_context_menu.add_command(label="Supprimer", command=self.delete_selected_client)
        
        # Bind context menu to clients tree
        self.clients_tree.bind("<Button-3>", self.show_clients_context_menu)  # Right click
        self.clients_tree.bind("<Double-1>", self.edit_selected_client)  # Double click to edit
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(list_frame, orient='vertical', command=self.clients_tree.yview)
        self.clients_tree.configure(yscrollcommand=scrollbar.set)
        
        self.clients_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        # Load clients
        self.refresh_clients_list()
    
    def setup_quotes_tab(self):
        """Setup quotes management tab"""
        # Title
        title_label = ttk.Label(self.quotes_frame, text="Gestion des Devis", font=('Arial', 16, 'bold'))
        title_label.pack(pady=10)
        
        # New quote button
        buttons_frame = ttk.Frame(self.quotes_frame)
        buttons_frame.pack(pady=5)
        
        new_quote_button = ttk.Button(buttons_frame, text="Nouveau Devis", command=self.new_quote)
        new_quote_button.pack(side='left', padx=(0, 10))
        
        delete_quote_button = ttk.Button(buttons_frame, text="Supprimer Devis", command=self.delete_selected_quote)
        delete_quote_button.pack(side='left')
        
        # Quotes list
        list_frame = ttk.LabelFrame(self.quotes_frame, text="Devis existants", padding=10)
        list_frame.pack(fill='both', expand=True, padx=10, pady=5)

        # Search toolbox for quotes
        search_frame = ttk.Frame(list_frame)
        search_frame.pack(fill='x', pady=(0, 10))

        ttk.Label(search_frame, text="Rechercher:").pack(side='left')
        self.quotes_search_var = tk.StringVar()
        quotes_search_entry = ttk.Entry(search_frame, textvariable=self.quotes_search_var)
        quotes_search_entry.pack(side='left', fill='x', expand=True, padx=(5, 5))
        ttk.Button(search_frame, text="Filtrer", command=self.apply_quotes_filter).pack(side='left')
        ttk.Button(search_frame, text="Réinitialiser", command=self.reset_quotes_filter).pack(side='left', padx=(5, 0))
        quotes_search_entry.bind('<Return>', lambda event: self.apply_quotes_filter())
        
        # Treeview for quotes list
        columns = ('Numéro', 'Client', 'Typologie', 'Date', 'Sites', 'Total HT', 'Total TTC', 'Facturé', 'Actions')
        self.quotes_tree = ttk.Treeview(list_frame, columns=columns, show='headings')
        
        for col in columns:
            self.quotes_tree.heading(col, text=col)
            if col == 'Facturé':
                self.quotes_tree.column(col, width=80)
            elif col == 'Sites':
                self.quotes_tree.column(col, width=120)
            elif col == 'Typologie':
                self.quotes_tree.column(col, width=80)
            else:
                self.quotes_tree.column(col, width=150)
        
        # Scrollbar
        scrollbar_quotes = ttk.Scrollbar(list_frame, orient='vertical', command=self.quotes_tree.yview)
        self.quotes_tree.configure(yscrollcommand=scrollbar_quotes.set)
        
        self.quotes_tree.pack(side='left', fill='both', expand=True)
        scrollbar_quotes.pack(side='right', fill='y')
        
        # Bind double-click to open quote
        self.quotes_tree.bind('<Double-1>', self.view_quote)
        
        # Bind single-click to handle Sites column clicks
        self.quotes_tree.bind('<Button-1>', self.on_quotes_tree_click)
        
        # Context menu for quotes
        self.quotes_context_menu = tk.Menu(self.root, tearoff=0)
        self.quotes_context_menu.add_command(label="Voir/Modifier", command=self.view_quote)
        self.quotes_context_menu.add_separator()
        self.quotes_context_menu.add_command(label="Convertir en Facture", command=self.convert_to_invoice)
        self.quotes_context_menu.add_command(label="Voir Facture Liée", command=self.view_linked_invoice)
        self.quotes_context_menu.add_separator()
        self.quotes_context_menu.add_command(label="Exporter PDF", command=self.export_quote_pdf)
        self.quotes_context_menu.add_command(label="Exporter Word", command=self.export_quote_word)
        self.quotes_context_menu.add_separator()
        self.quotes_context_menu.add_command(label="Supprimer", command=self.delete_selected_quote)
        
        self.quotes_tree.bind('<Button-3>', self.show_quotes_context_menu)
        
        # Load quotes
        self.refresh_quotes_list()
    
    def setup_invoices_tab(self):
        """Setup invoices management tab"""
        # Title
        title_label = ttk.Label(self.invoices_frame, text="Gestion des Factures", font=('Arial', 16, 'bold'))
        title_label.pack(pady=10)
        
        # Buttons frame for invoice actions
        invoices_buttons_frame = ttk.Frame(self.invoices_frame)
        invoices_buttons_frame.pack(pady=5)
        
        ttk.Button(invoices_buttons_frame, text="Supprimer Facture", 
                  command=self.delete_selected_invoice).pack(side='left')
        
        # Invoices list
        list_frame = ttk.LabelFrame(self.invoices_frame, text="Factures existantes", padding=10)
        list_frame.pack(fill='both', expand=True, padx=10, pady=5)

        # Search toolbox for invoices
        invoices_search_frame = ttk.Frame(list_frame)
        invoices_search_frame.pack(fill='x', pady=(0, 10))

        ttk.Label(invoices_search_frame, text="Rechercher:").pack(side='left')
        self.invoices_search_var = tk.StringVar()
        invoices_search_entry = ttk.Entry(invoices_search_frame, textvariable=self.invoices_search_var)
        invoices_search_entry.pack(side='left', fill='x', expand=True, padx=(5, 5))
        ttk.Button(invoices_search_frame, text="Filtrer", command=self.apply_invoices_filter).pack(side='left')
        ttk.Button(invoices_search_frame, text="Réinitialiser", command=self.reset_invoices_filter).pack(side='left', padx=(5, 0))
        invoices_search_entry.bind('<Return>', lambda event: self.apply_invoices_filter())
        
        # Treeview for invoices list
        columns = ('Numéro Facture', 'Bon de Commande', 'Client', 'Date', 'Sites', 'Total HT', 'Total TTC')
        self.invoices_tree = ttk.Treeview(list_frame, columns=columns, show='headings')
        
        for col in columns:
            self.invoices_tree.heading(col, text=col)
            if col == 'Sites':
                self.invoices_tree.column(col, width=120)
            else:
                self.invoices_tree.column(col, width=150)
        
        # Scrollbar
        scrollbar_invoices = ttk.Scrollbar(list_frame, orient='vertical', command=self.invoices_tree.yview)
        self.invoices_tree.configure(yscrollcommand=scrollbar_invoices.set)
        
        self.invoices_tree.pack(side='left', fill='both', expand=True)
        scrollbar_invoices.pack(side='right', fill='y')
        
        # Bind single-click to handle Sites column clicks
        self.invoices_tree.bind('<Button-1>', self.on_invoices_tree_click)
        
        # Context menu for invoices
        self.invoices_context_menu = tk.Menu(self.root, tearoff=0)
        self.invoices_context_menu.add_command(label="Voir", command=self.view_invoice)
        self.invoices_context_menu.add_command(label="Exporter PDF", command=self.export_invoice_pdf)
        self.invoices_context_menu.add_command(label="Exporter Word", command=self.export_invoice_word)
        self.invoices_context_menu.add_separator()
        self.invoices_context_menu.add_command(label="Supprimer", command=self.delete_selected_invoice)
        
        self.invoices_tree.bind('<Button-3>', self.show_invoices_context_menu)
        
        # Load invoices
        self.refresh_invoices_list()
    
    def add_client(self):
        """Add new client to database"""
        if not self.client_name_var.get().strip():
            messagebox.showerror("Erreur", "Le nom du client est obligatoire")
            return
        
        client = Client(
            name=self.client_name_var.get().strip(),
            siret=self.client_siret_var.get().strip(),
            address=self.client_address_text.get(1.0, tk.END).strip(),
            email=self.client_email_var.get().strip(),
            phone=self.client_phone_var.get().strip()
        )
        
        try:
            client_id = self.db.add_client(client)
            messagebox.showinfo("Succès", f"Client ajouté avec l'ID: {client_id}")
            
            # Clear form
            self.client_name_var.set("")
            self.client_siret_var.set("")
            self.client_address_text.delete(1.0, tk.END)
            self.client_email_var.set("")
            self.client_phone_var.set("")
            
            # Refresh list
            self.refresh_clients_list()
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'ajout du client: {str(e)}")
    
    def delete_selected_client(self):
        """Delete the selected client after simple confirmation"""
        selection = self.clients_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un client à supprimer")
            return
        
        # Get selected client info
        item = self.clients_tree.item(selection[0])
        client_id = int(item['values'][0])
        client_name = item['values'][1]
        
        # Simple confirmation
        result = messagebox.askyesno(
            "Confirmer la suppression",
            f"Êtes-vous sûr de vouloir supprimer le client '{client_name}' ?\n\n"
            "Cette action est irréversible.",
            icon='question'
        )
        
        if not result:
            return
        
        try:
            success = self.db.delete_client(client_id)
            if success:
                messagebox.showinfo("Succès", f"Client '{client_name}' supprimé avec succès")
                self.refresh_clients_list()
            
        except ValueError as e:
            messagebox.showerror("Suppression impossible", str(e))
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la suppression: {str(e)}")
    
    def edit_selected_client(self):
        """Edit the selected client (placeholder for future implementation)"""
        selection = self.clients_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un client à modifier")
            return
        
        item = self.clients_tree.item(selection[0])
        client_id = int(item['values'][0])
        client_name = item['values'][1]
        
        # TODO: Implement client editing dialog
        messagebox.showinfo("Fonctionnalité à venir", 
                           f"Modification du client '{client_name}' (ID: {client_id})\n\n"
                           "Cette fonctionnalité sera implémentée prochainement.")
    
    def show_clients_context_menu(self, event):
        """Show context menu for clients tree"""
        # Select item under cursor
        item = self.clients_tree.identify_row(event.y)
        if item:
            self.clients_tree.selection_set(item)
            self.clients_context_menu.post(event.x_root, event.y_root)
    
    def delete_selected_quote(self):
        """Delete the selected quote after simple confirmation"""
        selection = self.quotes_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un devis à supprimer")
            return
        
        # Get selected quote info
        item = self.quotes_tree.item(selection[0])
        quote_id = int(item['tags'][0])  # ID is stored in tags
        quote_number = item['values'][0]  # Quote number is first column
        
        # Simple confirmation
        result = messagebox.askyesno(
            "Confirmer la suppression",
            f"Êtes-vous sûr de vouloir supprimer le devis '{quote_number}' ?\n\n"
            "Cette action est irréversible.",
            icon='question'
        )
        
        if not result:
            return
        
        try:
            success = self.db.delete_quote(quote_id)
            if success:
                messagebox.showinfo("Succès", f"Devis '{quote_number}' supprimé avec succès")
                self.refresh_quotes_list()
                # Also refresh invoices list in case this was linked to an invoice
                self.refresh_invoices_list()
            
        except ValueError as e:
            messagebox.showerror("Suppression impossible", str(e))
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la suppression: {str(e)}")
    
    def delete_selected_invoice(self):
        """Delete the selected invoice after simple confirmation"""
        selection = self.invoices_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner une facture à supprimer")
            return
        
        # Get selected invoice info
        item = self.invoices_tree.item(selection[0])
        invoice_id = int(item['tags'][0])  # ID is stored in tags
        invoice_number = item['values'][0]  # Invoice number is first column
        
        # Simple confirmation
        result = messagebox.askyesno(
            "Confirmer la suppression",
            f"Êtes-vous sûr de vouloir supprimer la facture '{invoice_number}' ?\n\n"
            "Cette action est irréversible et remettra le devis original en statut non facturé.",
            icon='question'
        )
        
        if not result:
            return
        
        try:
            success = self.db.delete_quote(invoice_id)  # Invoices are stored in the same table
            if success:
                messagebox.showinfo("Succès", f"Facture '{invoice_number}' supprimée avec succès")
                self.refresh_invoices_list()
                # Also refresh quotes list to update invoice status
                self.refresh_quotes_list()
            
        except ValueError as e:
            messagebox.showerror("Suppression impossible", str(e))
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la suppression: {str(e)}")
    
    def refresh_clients_list(self):
        """Refresh the clients list display"""
        # Clear existing items
        for item in self.clients_tree.get_children():
            self.clients_tree.delete(item)
        
        # Load clients from database
        clients = self.db.get_clients()
        for client in clients:
            self.clients_tree.insert('', 'end', values=(
                client.id,
                client.name,
                client.siret or "",
                client.email or "",
                client.phone or ""
            ))
    
    def refresh_quotes_list(self):
        """Refresh the quotes list display"""
        # Clear existing items
        for item in self.quotes_tree.get_children():
            self.quotes_tree.delete(item)
        
        # Load quotes from database
        quotes = self.db.get_quotes(is_invoice=False)
        search_text = ""
        if hasattr(self, 'quotes_search_var'):
            search_text = self.quotes_search_var.get().strip().lower()
        linked_invoices = {}
        if search_text:
            linked_invoices = {inv.id: inv for inv in self.db.get_quotes(is_invoice=True)}
        for quote in quotes:
            client_name = quote.client.name if quote.client else "Client inconnu"
            
            # Use quote_date if available, otherwise fall back to created_at
            if quote.quote_date:
                date_str = quote.quote_date.strftime("%d/%m/%Y")
            elif quote.created_at:
                date_str = quote.created_at.strftime("%d/%m/%Y")
            else:
                date_str = ""
            
            # Determine invoice status and action text
            if quote.is_invoiced:
                invoice_status = "✓ Facturé"
                # Get the actual invoice number for display
                if quote.linked_invoice_id:
                    # Get the linked invoice to show its number
                    invoices = self.db.get_quotes(is_invoice=True)
                    linked_invoice = next((inv for inv in invoices if inv.id == quote.linked_invoice_id), None)
                    if linked_invoice and linked_invoice.invoice_number:
                        action_text = f"Voir {linked_invoice.invoice_number}"
                    else:
                        action_text = f"Voir facture #{quote.linked_invoice_id}"
                else:
                    action_text = "Facturé"
            else:
                invoice_status = "Non facturé"
                action_text = "Actions"
            
            if search_text:
                linked_invoice_number = ""
                if quote.is_invoiced and quote.linked_invoice_id:
                    linked_invoice = linked_invoices.get(quote.linked_invoice_id)
                    if linked_invoice and linked_invoice.invoice_number:
                        linked_invoice_number = linked_invoice.invoice_number
                searchable_fields = [
                    quote.number or "",
                    client_name,
                    quote.typology or "",
                    date_str,
                    quote.site_numbers_display or "",
                    invoice_status,
                    action_text,
                    linked_invoice_number,
                ]
                if not any(search_text in str(field).lower() for field in searchable_fields if field):
                    continue
            
            item_id = self.quotes_tree.insert('', 'end', values=(
                quote.number,
                client_name,
                quote.typology or "",  # Add typology display
                date_str,
                quote.site_numbers_display,  # Show site numbers
                f"{quote.total_ht:.2f} €",
                f"{quote.total_ttc:.2f} €",
                invoice_status,
                action_text
            ), tags=(str(quote.id), 'invoiced' if quote.is_invoiced else 'not_invoiced'))
            
        # Configure styling for invoiced and non-invoiced items
        self.quotes_tree.tag_configure('invoiced', background='#e8f5e8')
        self.quotes_tree.tag_configure('not_invoiced', background='white')
    
    def refresh_invoices_list(self):
        """Refresh the invoices list display"""
        # Clear existing items
        for item in self.invoices_tree.get_children():
            self.invoices_tree.delete(item)
        
        # Load invoices from database
        invoices = self.db.get_quotes(is_invoice=True)
        search_text = ""
        if hasattr(self, 'invoices_search_var'):
            search_text = self.invoices_search_var.get().strip().lower()
        for invoice in invoices:
            client_name = invoice.client.name if invoice.client else "Client inconnu"
            
            # Use intervention_date if available for invoices, otherwise fall back to created_at
            if invoice.intervention_date:
                date_str = invoice.intervention_date.strftime("%d/%m/%Y")
            elif invoice.created_at:
                date_str = invoice.created_at.strftime("%d/%m/%Y")
            else:
                date_str = ""
            
            if search_text:
                searchable_fields = [
                    invoice.invoice_number or "",
                    invoice.order_number or "",
                    client_name,
                    date_str,
                    invoice.site_numbers_display or "",
                    f"{invoice.total_ht:.2f}",
                    f"{invoice.total_ttc:.2f}",
                ]
                if not any(search_text in str(field).lower() for field in searchable_fields if field):
                    continue
            
            self.invoices_tree.insert('', 'end', values=(
                invoice.invoice_number or "N/A",
                invoice.order_number or "N/A",
                client_name,
                date_str,
                invoice.site_numbers_display,  # Show site numbers
                f"{invoice.total_ht:.2f} €",
                f"{invoice.total_ttc:.2f} €"
            ), tags=(str(invoice.id),))
    
    def apply_quotes_filter(self, event=None):
        """Apply search filter to quotes"""
        self.refresh_quotes_list()

    def reset_quotes_filter(self):
        """Reset quotes search filter"""
        if hasattr(self, 'quotes_search_var'):
            self.quotes_search_var.set("")
        self.refresh_quotes_list()

    def apply_invoices_filter(self, event=None):
        """Apply search filter to invoices"""
        self.refresh_invoices_list()

    def reset_invoices_filter(self):
        """Reset invoices search filter"""
        if hasattr(self, 'invoices_search_var'):
            self.invoices_search_var.set("")
        self.refresh_invoices_list()

    def new_quote(self):
        """Open new quote dialog"""
        QuoteDialog(self.root, self.db, callback=self.refresh_quotes_list)
    
    def on_quotes_tree_click(self, event):
        """Handle clicks on the quotes tree, specifically for Sites column"""
        # Identify what was clicked
        region = self.quotes_tree.identify("region", event.x, event.y)
        if region == "cell":
            # Get the column that was clicked using the correct Tkinter method
            column = self.quotes_tree.identify("column", event.x, event.y)
            # Column #4 is the Sites column (0-indexed: Numéro, Client, Date, Sites)
            if column == '#4':  # Sites column
                # Get the selected item
                item = self.quotes_tree.identify("item", event.x, event.y)
                if item:
                    self.show_sites_details(item, is_invoice=False)
    
    def show_sites_details(self, item_id, is_invoice=False):
        """Show detailed view of all sites for the selected quote/invoice"""
        try:
            item = self.quotes_tree.item(item_id) if not is_invoice else self.invoices_tree.item(item_id)
            # Extract quote ID from the first tag (tags[0] is the ID)
            quote_id = int(item['tags'][0])
            
            # Get quote/invoice from database
            quotes = self.db.get_quotes(is_invoice=is_invoice)
            selected_quote = next((q for q in quotes if q.id == quote_id), None)
            
            if not selected_quote:
                messagebox.showerror("Erreur", "Document non trouvé")
                return
            
            # Create popup window
            popup = tk.Toplevel(self.root)
            popup.title(f"Sites - {'Facture' if is_invoice else 'Devis'} {selected_quote.invoice_number if is_invoice else selected_quote.number}")
            popup.geometry("800x400")
            popup.transient(self.root)
            popup.grab_set()
            
            # Main frame
            main_frame = ttk.Frame(popup, padding=10)
            main_frame.pack(fill='both', expand=True)
            
            # Title
            title_text = f"Sites d'intervention - {'Facture' if is_invoice else 'Devis'} {selected_quote.invoice_number if is_invoice else selected_quote.number}"
            title_label = ttk.Label(main_frame, text=title_text, font=('Arial', 12, 'bold'))
            title_label.pack(pady=(0, 10))
            
            # Client info
            client_info = f"Client: {selected_quote.client.name if selected_quote.client else 'Inconnu'}"
            client_label = ttk.Label(main_frame, text=client_info)
            client_label.pack(anchor='w', pady=(0, 10))
            
            # Sites table
            columns = ('N° Site', 'Adresse', 'Description', 'Prix HT')
            sites_tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=12)
            
            for col in columns:
                sites_tree.heading(col, text=col)
                if col == 'N° Site':
                    sites_tree.column(col, width=120)
                elif col == 'Adresse':
                    sites_tree.column(col, width=200)
                elif col == 'Description':
                    sites_tree.column(col, width=300)
                elif col == 'Prix HT':
                    sites_tree.column(col, width=100)
            
            # Add scrollbar
            scrollbar = ttk.Scrollbar(main_frame, orient='vertical', command=sites_tree.yview)
            sites_tree.configure(yscrollcommand=scrollbar.set)
            
            # Pack tree and scrollbar
            sites_tree.pack(side='left', fill='both', expand=True)
            scrollbar.pack(side='right', fill='y')
            
            # Populate with sites data
            total_ht = 0
            site_count = 0
            
            # Add sites from the new multi-site system
            for site in selected_quote.sites:
                sites_tree.insert('', 'end', values=(
                    site.site_number,
                    site.site_address,
                    site.description,
                    f"{site.price_ht:.2f} €"
                ))
                total_ht += site.price_ht
                site_count += 1
            
            # Add legacy single site if exists and not already included
            if selected_quote.site_number and selected_quote.site_number not in [s.site_number for s in selected_quote.sites]:
                # For legacy items, we need to get the descriptions from quote_items
                descriptions = []
                for item in selected_quote.items:
                    descriptions.append(item.description)
                legacy_description = " | ".join(descriptions) if descriptions else "Service"
                
                legacy_total = sum(item.total_ht for item in selected_quote.items)
                sites_tree.insert('', 'end', values=(
                    selected_quote.site_number,
                    selected_quote.site_address or "",
                    legacy_description,
                    f"{legacy_total:.2f} €"
                ))
                total_ht += legacy_total
                site_count += 1
            
            # If no sites found, show message
            if site_count == 0:
                sites_tree.insert('', 'end', values=(
                    "Aucun site",
                    "",
                    "Aucune information de site disponible",
                    "0.00 €"
                ))
            
            # Total info frame
            totals_frame = ttk.Frame(popup)
            totals_frame.pack(fill='x', padx=10, pady=10)
            
            total_tva = total_ht * BUSINESS_CONFIG.get('default_vat_rate', 0.20)
            total_ttc = total_ht + total_tva
            
            totals_text = f"Total: {site_count} site(s) | HT: {total_ht:.2f} € | TVA: {total_tva:.2f} € | TTC: {total_ttc:.2f} €"
            totals_label = ttk.Label(totals_frame, text=totals_text, font=('Arial', 10, 'bold'))
            totals_label.pack()
            
            # Close button
            close_button = ttk.Button(totals_frame, text="Fermer", command=popup.destroy)
            close_button.pack(pady=(10, 0))
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'affichage des sites: {str(e)}")
    
    def view_quote(self, event=None):
        """View/edit selected quote"""
        selection = self.quotes_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un devis à modifier")
            return
        
        try:
            item = self.quotes_tree.item(selection[0])
            # Extract quote ID from the first tag (tags[0] is the ID)
            quote_id = int(item['tags'][0])
            
            # Get the quote from database
            quotes = self.db.get_quotes(is_invoice=False)
            selected_quote = next((q for q in quotes if q.id == quote_id), None)
            
            if not selected_quote:
                messagebox.showerror("Erreur", "Devis introuvable")
                return
            
            # Open quote dialog in edit mode
            dialog = QuoteDialog(
                parent=self.root, 
                db=self.db, 
                quote=selected_quote,  # Pass the existing quote for editing
                callback=self.refresh_quotes_list  # Refresh list after editing
            )
            
        except ValueError:
            messagebox.showerror("Erreur", "ID de devis invalide")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'ouverture du devis: {str(e)}")
    
    def view_invoice(self, event=None):
        """View selected invoice"""
        selection = self.invoices_tree.selection()
        if selection:
            item = self.invoices_tree.item(selection[0])
            invoice_id = int(item['tags'][0])
            messagebox.showinfo("Invoice", f"Viewing invoice ID: {invoice_id}")
    
    def on_invoices_tree_click(self, event):
        """Handle clicks on the invoices tree, specifically for Sites column"""
        # Identify what was clicked
        region = self.invoices_tree.identify("region", event.x, event.y)
        if region == "cell":
            # Get the column that was clicked using the correct Tkinter method
            column = self.invoices_tree.identify("column", event.x, event.y)
            # Column #5 is the Sites column for invoices (0-indexed: Numéro Facture, Bon de Commande, Client, Date, Sites)
            if column == '#5':  # Sites column
                # Get the selected item
                item = self.invoices_tree.identify("item", event.x, event.y)
                if item:
                    self.show_sites_details(item, is_invoice=True)
    
    def convert_to_invoice(self):
        """Convert selected quote to invoice"""
        selection = self.quotes_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un devis")
            return
        
        item = self.quotes_tree.item(selection[0])
        # Extract quote ID from the first tag (tags[0] is the ID)
        quote_id = int(item['tags'][0])
        
        # Show conversion dialog
        dialog = ConvertToInvoiceDialog(self.root)
        self.root.wait_window(dialog.dialog)
        
        if dialog.result:
            try:
                invoice_number = self.db.convert_quote_to_invoice(
                    quote_id, 
                    dialog.result['order_number'],
                    dialog.result['intervention_date']
                )
                messagebox.showinfo("Succès", f"Devis converti en facture: {invoice_number}")
                self.refresh_quotes_list()
                self.refresh_invoices_list()
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de la conversion: {str(e)}")
    
    def export_quote_pdf(self):
        """Export selected quote as PDF"""
        selection = self.quotes_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un devis")
            return
        
        self._export_document_pdf(selection[0], False)
    
    def export_invoice_pdf(self):
        """Export selected invoice as PDF"""
        selection = self.invoices_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner une facture")
            return
        
        self._export_document_pdf(selection[0], True)
    
    def export_quote_word(self):
        """Export selected quote as Word document"""
        selection = self.quotes_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un devis")
            return
        
        self._export_document_word(selection[0], False)
    
    def export_invoice_word(self):
        """Export selected invoice as Word document"""
        selection = self.invoices_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner une facture")
            return
        
        self._export_document_word(selection[0], True)
    
    def _export_document_pdf(self, selection, is_invoice):
        """Export document as PDF"""
        try:
            if not PDF_AVAILABLE:
                messagebox.showerror("Erreur", "reportlab n'est pas installé. Export PDF indisponible.")
                return
            
            # Get quote/invoice from database
            tree = self.invoices_tree if is_invoice else self.quotes_tree
            item = tree.item(selection)
            doc_id = int(item['tags'][0])
            
            quotes = self.db.get_quotes(is_invoice=is_invoice)
            quote = next((q for q in quotes if q.id == doc_id), None)
            
            if not quote:
                messagebox.showerror("Erreur", "Document non trouvé")
                return
            
            # Ask for save location
            doc_type = "facture" if is_invoice else "devis"
            filename = f"{doc_type}_{quote.number if not is_invoice else quote.invoice_number}.pdf"
            filepath = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")],
                initialfile=filename
            )
            
            if filepath:
                PDFGenerator.generate_quote_pdf(quote, filepath)
                messagebox.showinfo("Succès", f"PDF exporté: {filepath}")
                
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'export PDF: {str(e)}")
    
    def _export_document_word(self, selection, is_invoice):
        """Export document as Word"""
        try:
            if not DOCX_AVAILABLE:
                messagebox.showerror("Erreur", "python-docx n'est pas installé. Export Word indisponible.")
                return
            
            # Get quote/invoice from database
            tree = self.invoices_tree if is_invoice else self.quotes_tree
            item = tree.item(selection)
            doc_id = int(item['tags'][0])
            
            quotes = self.db.get_quotes(is_invoice=is_invoice)
            quote = next((q for q in quotes if q.id == doc_id), None)
            
            if not quote:
                messagebox.showerror("Erreur", "Document non trouvé")
                return
            
            # Ask for save location
            doc_type = "facture" if is_invoice else "devis"
            filename = f"{doc_type}_{quote.number if not is_invoice else quote.invoice_number}.docx"
            filepath = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")],
                initialfile=filename
            )
            
            if filepath:
                WordGenerator.generate_quote_docx(quote, filepath)
                messagebox.showinfo("Succès", f"Document Word exporté: {filepath}")
                
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'export Word: {str(e)}")
    
    def show_quotes_context_menu(self, event):
        """Show context menu for quotes with dynamic options based on invoice status"""
        selection = self.quotes_tree.selection()
        if not selection:
            return
            
        item = self.quotes_tree.item(selection[0])
        # Extract quote ID from the first tag (tags[0] is the ID)
        # Extract quote ID from the first tag (tags[0] is the ID)
        quote_id = int(item['tags'][0])
        
        # Get quote to check if it's invoiced
        quotes = self.db.get_quotes(is_invoice=False)
        current_quote = next((q for q in quotes if q.id == quote_id), None)
        
        # Clear existing menu
        self.quotes_context_menu.delete(0, 'end')
        
        # Add standard options
        self.quotes_context_menu.add_command(label="Voir/Modifier", command=self.view_quote)
        self.quotes_context_menu.add_separator()
        
        if current_quote and current_quote.is_invoiced:
            # Quote has been invoiced - show link to invoice instead of convert option
            self.quotes_context_menu.add_command(label="Voir Facture Liée", command=self.view_linked_invoice)
            self.quotes_context_menu.add_command(label="Aller à la Facture", command=self.go_to_linked_invoice)
        else:
            # Quote not invoiced - show convert option
            self.quotes_context_menu.add_command(label="Convertir en Facture", command=self.convert_to_invoice)
        
        self.quotes_context_menu.add_separator()
        self.quotes_context_menu.add_command(label="Exporter PDF", command=self.export_quote_pdf)
        self.quotes_context_menu.add_command(label="Exporter Word", command=self.export_quote_word)
        
        try:
            self.quotes_context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.quotes_context_menu.grab_release()
    
    def view_linked_invoice(self):
        """View the invoice linked to the selected quote"""
        selection = self.quotes_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un devis")
            return
        
        item = self.quotes_tree.item(selection[0])
        # Extract quote ID from the first tag (tags[0] is the ID)
        quote_id = int(item['tags'][0])
        
        # Get quote to find linked invoice
        quotes = self.db.get_quotes(is_invoice=False)
        quote = next((q for q in quotes if q.id == quote_id), None)
        
        if not quote or not quote.is_invoiced or not quote.linked_invoice_id:
            messagebox.showinfo("Information", "Ce devis n'a pas de facture liée")
            return
        
        # Get invoice details
        invoices = self.db.get_quotes(is_invoice=True)
        linked_invoice = next((inv for inv in invoices if inv.id == quote.linked_invoice_id), None)
        
        if linked_invoice:
            messagebox.showinfo("Facture Liée", 
                f"Facture: {linked_invoice.invoice_number}\n"
                f"Bon de commande: {linked_invoice.order_number}\n"
                f"Date: {linked_invoice.created_at.strftime('%d/%m/%Y') if linked_invoice.created_at else 'N/A'}\n"
                f"Montant TTC: {linked_invoice.total_ttc:.2f} €")
        else:
            messagebox.showerror("Erreur", "Facture liée introuvable")
    
    def go_to_linked_invoice(self):
        """Navigate to the invoices tab and highlight the linked invoice"""
        selection = self.quotes_tree.selection()
        if not selection:
            return
        
        item = self.quotes_tree.item(selection[0])
        # Extract quote ID from the first tag (tags[0] is the ID)
        quote_id = int(item['tags'][0])
        
        # Get quote to find linked invoice
        quotes = self.db.get_quotes(is_invoice=False)
        quote = next((q for q in quotes if q.id == quote_id), None)
        
        if not quote or not quote.is_invoiced or not quote.linked_invoice_id:
            messagebox.showinfo("Information", "Ce devis n'a pas de facture liée")
            return
        
        # Switch to invoices tab
        # Find the notebook widget and switch to invoices tab
        for widget in self.root.winfo_children():
            if isinstance(widget, ttk.Notebook):
                widget.select(2)  # Invoices tab is the third tab (index 2)
                break
        
        # Refresh invoices list and try to highlight the linked invoice
        self.refresh_invoices_list()
        
        # Find and select the linked invoice in the invoices tree
        for item_id in self.invoices_tree.get_children():
            item = self.invoices_tree.item(item_id)
            if item['tags'] and int(item['tags'][0]) == quote.linked_invoice_id:
                self.invoices_tree.selection_set(item_id)
                self.invoices_tree.see(item_id)
                break
    
    def show_invoices_context_menu(self, event):
        """Show context menu for invoices"""
        try:
            self.invoices_context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.invoices_context_menu.grab_release()
    
    def run(self):
        """Run the application"""
        self.root.mainloop()

class ConvertToInvoiceDialog:
    """Dialog for converting quote to invoice with order number and intervention date"""
    
    def __init__(self, parent):
        self.result = None
        
        # Create dialog window
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Convertir en Facture")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        self.setup_ui()
        
        # Size dialog based on content and center it on screen
        self.dialog.update_idletasks()
        width = max(420, self.dialog.winfo_reqwidth() + 40)
        height = max(320, self.dialog.winfo_reqheight() + 40)
        x = (self.dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (self.dialog.winfo_screenheight() // 2) - (height // 2)
        self.dialog.geometry(f"{width}x{height}+{x}+{y}")
        self.dialog.resizable(False, False)
    
    def setup_ui(self):
        """Setup dialog UI"""
        main_frame = ttk.Frame(self.dialog, padding=20)
        main_frame.pack(fill='both', expand=True)
        
        # Title
        title_label = ttk.Label(main_frame, text="Conversion Devis → Facture", font=('Arial', 12, 'bold'))
        title_label.pack(pady=(0, 20))
        
        # Order number frame
        order_frame = ttk.LabelFrame(main_frame, text="Bon de commande", padding=10)
        order_frame.pack(fill='x', pady=(0, 15))
        
        ttk.Label(order_frame, text="Numéro du bon de commande:").pack(anchor='w')
        self.order_var = tk.StringVar()
        ttk.Entry(order_frame, textvariable=self.order_var, width=30).pack(fill='x', pady=(5, 0))
        
        # Intervention date frame
        date_frame = ttk.LabelFrame(main_frame, text="Date d'intervention", padding=10)
        date_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(date_frame, text="Date d'intervention:").pack(anchor='w')
        
        # Date entry fields
        self.day_var = tk.StringVar()
        self.month_var = tk.StringVar()
        self.year_var = tk.StringVar()
        
        # Set default date to today
        today = datetime.date.today()
        self.day_var.set(f"{today.day:02d}")
        self.month_var.set(f"{today.month:02d}")
        self.year_var.set(str(today.year))
        
        date_entry_frame = ttk.Frame(date_frame)
        date_entry_frame.pack(pady=(5, 0))
        
        ttk.Entry(date_entry_frame, textvariable=self.day_var, width=3).pack(side='left')
        ttk.Label(date_entry_frame, text="/").pack(side='left')
        ttk.Entry(date_entry_frame, textvariable=self.month_var, width=3).pack(side='left')
        ttk.Label(date_entry_frame, text="/").pack(side='left')
        ttk.Entry(date_entry_frame, textvariable=self.year_var, width=5).pack(side='left')
        ttk.Label(date_entry_frame, text="(JJ/MM/AAAA)").pack(side='left', padx=(10, 0))
        
        # Buttons
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill='x')
        
        ttk.Button(buttons_frame, text="Annuler", command=self.cancel).pack(side='right', padx=(10, 0))
        ttk.Button(buttons_frame, text="Convertir", command=self.convert).pack(side='right')
    
    def convert(self):
        """Validate inputs and close with result"""
        order_number = self.order_var.get().strip()
        if not order_number:
            messagebox.showerror("Erreur", "Le numéro de bon de commande est obligatoire")
            return
        
        # Validate date
        try:
            day = int(self.day_var.get())
            month = int(self.month_var.get())
            year = int(self.year_var.get())
            intervention_date = datetime.date(year, month, day)
        except (ValueError, TypeError):
            messagebox.showerror("Erreur", "Date d'intervention invalide. Veuillez vérifier le format JJ/MM/AAAA")
            return
        
        self.result = {
            'order_number': order_number,
            'intervention_date': intervention_date
        }
        self.dialog.destroy()
    
    def cancel(self):
        """Cancel dialog"""
        self.result = None
        self.dialog.destroy()

class QuoteDialog:
    """Dialog for creating/editing quotes"""
    
    def __init__(self, parent, db: DatabaseManager, quote: Quote = None, callback=None):
        self.db = db
        self.callback = callback
        self.quote = quote or Quote()
        
        # Site editing state variables
        self.editing_site_index = None  # Index of the site being edited
        self.is_editing_site = False    # Flag to track if we're in edit mode
        
        # Create dialog window
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Nouveau Devis" if not quote else "Modifier Devis")
        self.dialog.geometry("900x750")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        self.setup_ui()
        
        if quote:
            self.load_quote_data()
    
    def setup_ui(self):
        """Setup dialog UI"""
        # Main frame
        main_frame = ttk.Frame(self.dialog, padding=10)
        main_frame.pack(fill='both', expand=True)
        
        # Client selection
        client_frame = ttk.LabelFrame(main_frame, text="Client", padding=10)
        client_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(client_frame, text="Sélectionner un client:").pack(anchor='w')
        
        # Client combobox
        self.client_var = tk.StringVar()
        self.client_combo = ttk.Combobox(client_frame, textvariable=self.client_var, width=50)
        self.client_combo.pack(fill='x', pady=(5, 10))
        
        # Date selection
        date_frame = ttk.Frame(client_frame)
        date_frame.pack(fill='x')
        
        ttk.Label(date_frame, text="Date du devis:").pack(side='left')
        
        # Date entry fields
        self.quote_day_var = tk.StringVar()
        self.quote_month_var = tk.StringVar()
        self.quote_year_var = tk.StringVar()
        
        # Set default date to today
        today = datetime.date.today()
        self.quote_day_var.set(f"{today.day:02d}")
        self.quote_month_var.set(f"{today.month:02d}")
        self.quote_year_var.set(str(today.year))
        
        date_entry_frame = ttk.Frame(date_frame)
        date_entry_frame.pack(side='left', padx=(10, 0))
        
        ttk.Entry(date_entry_frame, textvariable=self.quote_day_var, width=3).pack(side='left')
        ttk.Label(date_entry_frame, text="/").pack(side='left')
        ttk.Entry(date_entry_frame, textvariable=self.quote_month_var, width=3).pack(side='left')
        ttk.Label(date_entry_frame, text="/").pack(side='left')
        ttk.Entry(date_entry_frame, textvariable=self.quote_year_var, width=5).pack(side='left')
        ttk.Label(date_entry_frame, text="(JJ/MM/AAAA)").pack(side='left', padx=(5, 0))
        
        # Typology selection
        typology_frame = ttk.Frame(client_frame)
        typology_frame.pack(fill='x', pady=(10, 0))
        
        ttk.Label(typology_frame, text="Typologie:").pack(side='left')
        
        self.typology_var = tk.StringVar()
        typology_values = ["", "Pylonne", "TT", "CDE", "Eglise", "Autre"]  # Empty option first
        typology_combo = ttk.Combobox(typology_frame, textvariable=self.typology_var, 
                                     values=typology_values, state="readonly", width=15)
        typology_combo.pack(side='left', padx=(10, 0))
        
        # Load clients
        self.load_clients()
        
        # Sites management frame
        sites_frame = ttk.LabelFrame(main_frame, text="Sites d'Intervention", padding=10)
        sites_frame.pack(fill='both', expand=True, pady=(0, 10))
        
        # Add site form
        add_site_frame = ttk.Frame(sites_frame)
        add_site_frame.pack(fill='x', pady=(0, 10))
        
        # Site form fields - Multiple rows for detailed address
        site_fields_frame = ttk.Frame(add_site_frame)
        site_fields_frame.pack(fill='x')
        
        # Row 1: Site number and Address
        ttk.Label(site_fields_frame, text="N° Site:").grid(row=0, column=0, sticky='w', padx=(0, 5))
        self.site_number_var = tk.StringVar()
        ttk.Entry(site_fields_frame, textvariable=self.site_number_var, width=15).grid(row=0, column=1, sticky='ew', padx=(0, 10))
        
        ttk.Label(site_fields_frame, text="Adresse:").grid(row=0, column=2, sticky='w', padx=(0, 5))
        self.site_address_var = tk.StringVar()
        ttk.Entry(site_fields_frame, textvariable=self.site_address_var, width=40).grid(row=0, column=3, columnspan=2, sticky='ew', padx=(0, 10))
        
        # Row 2: Postal code and City
        ttk.Label(site_fields_frame, text="Code postal:").grid(row=1, column=0, sticky='w', padx=(0, 5), pady=(5, 0))
        self.site_postal_code_var = tk.StringVar()
        ttk.Entry(site_fields_frame, textvariable=self.site_postal_code_var, width=15).grid(row=1, column=1, sticky='ew', padx=(0, 10), pady=(5, 0))
        
        ttk.Label(site_fields_frame, text="Ville:").grid(row=1, column=2, sticky='w', padx=(0, 5), pady=(5, 0))
        self.site_city_var = tk.StringVar()
        ttk.Entry(site_fields_frame, textvariable=self.site_city_var, width=30).grid(row=1, column=3, columnspan=2, sticky='ew', padx=(0, 10), pady=(5, 0))
        
        # Row 3: Latitude and Longitude  
        ttk.Label(site_fields_frame, text="Latitude:").grid(row=2, column=0, sticky='w', padx=(0, 5), pady=(5, 0))
        self.site_latitude_var = tk.StringVar()
        ttk.Entry(site_fields_frame, textvariable=self.site_latitude_var, width=15).grid(row=2, column=1, sticky='ew', padx=(0, 10), pady=(5, 0))
        
        ttk.Label(site_fields_frame, text="Longitude:").grid(row=2, column=2, sticky='w', padx=(0, 5), pady=(5, 0))
        self.site_longitude_var = tk.StringVar()
        ttk.Entry(site_fields_frame, textvariable=self.site_longitude_var, width=15).grid(row=2, column=3, sticky='ew', padx=(0, 10), pady=(5, 0))
        
        # Row 4: Description and Price
        ttk.Label(site_fields_frame, text="Description:").grid(row=3, column=0, sticky='w', padx=(0, 5), pady=(5, 0))
        self.site_description_var = tk.StringVar()
        ttk.Entry(site_fields_frame, textvariable=self.site_description_var, width=50).grid(row=3, column=1, columnspan=3, sticky='ew', padx=(0, 10), pady=(5, 0))
        
        ttk.Label(site_fields_frame, text="Prix HT:").grid(row=3, column=4, sticky='w', padx=(0, 5), pady=(5, 0))
        self.site_price_var = tk.StringVar()
        ttk.Entry(site_fields_frame, textvariable=self.site_price_var, width=15).grid(row=3, column=5, sticky='ew', pady=(5, 0))
        
        # Configure grid weights
        site_fields_frame.columnconfigure(1, weight=1)
        site_fields_frame.columnconfigure(3, weight=2)
        
        # Add site button on its own row for better positioning
        button_frame = ttk.Frame(add_site_frame)
        button_frame.pack(fill='x', pady=(10, 0))
        
        # Store button reference for dynamic text changes
        self.add_site_button = ttk.Button(button_frame, text="Ajouter Site", command=self.add_or_update_site)
        self.add_site_button.pack(side='right')
        
        # Cancel edit button (initially hidden)
        self.cancel_edit_button = ttk.Button(button_frame, text="Annuler Modification", command=self.cancel_edit_site)
        self.cancel_edit_button.pack(side='right', padx=(0, 5))
        self.cancel_edit_button.pack_forget()  # Hide initially
        
        # Sites list - Updated columns for detailed address
        columns = ('N° Site', 'Adresse Complète', 'Coordonnées', 'Description', 'Prix HT')
        
        # Create a container frame for the tree and buttons
        tree_container = ttk.Frame(sites_frame)
        tree_container.pack(fill='both', expand=True)
        
        # Tree and scrollbar frame
        tree_frame = ttk.Frame(tree_container)
        tree_frame.pack(fill='both', expand=True)
        
        self.sites_tree = ttk.Treeview(tree_frame, columns=columns, show='headings', height=6)
        
        for col in columns:
            self.sites_tree.heading(col, text=col)
            if col == 'N° Site':
                self.sites_tree.column(col, width=80)
            elif col == 'Adresse Complète':
                self.sites_tree.column(col, width=200)
            elif col == 'Coordonnées':
                self.sites_tree.column(col, width=150)
            elif col == 'Description':
                self.sites_tree.column(col, width=250)
            elif col == 'Prix HT':
                self.sites_tree.column(col, width=80)
        
        # Scrollbar for sites
        sites_scrollbar = ttk.Scrollbar(tree_frame, orient='vertical', command=self.sites_tree.yview)
        self.sites_tree.configure(yscrollcommand=sites_scrollbar.set)
        
        self.sites_tree.pack(side='left', fill='both', expand=True)
        sites_scrollbar.pack(side='right', fill='y')
        
        # Site action buttons frame - ALWAYS VISIBLE
        site_buttons_frame = ttk.Frame(tree_container)
        site_buttons_frame.pack(fill='x', pady=(5, 0))
        
        # Edit site button
        edit_site_button = ttk.Button(site_buttons_frame, text="Modifier Site", command=self.edit_site)
        edit_site_button.pack(side='left', padx=(0, 5))
        
        # Remove site button
        remove_site_button = ttk.Button(site_buttons_frame, text="Supprimer Site", command=self.remove_site)
        remove_site_button.pack(side='left')
        
        # Totals frame
        totals_frame = ttk.LabelFrame(main_frame, text="Totaux", padding=10)
        totals_frame.pack(fill='x', pady=(0, 10))
        
        self.totals_label = ttk.Label(totals_frame, text="Total HT: 0.00 € | TVA: 0.00 € | TTC: 0.00 €")
        self.totals_label.pack()
        
        # Buttons frame
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill='x', pady=(10, 0))
        
        # Add some visual separation before buttons
        separator = ttk.Separator(buttons_frame, orient='horizontal')
        separator.pack(fill='x', pady=(0, 10))
        
        ttk.Button(buttons_frame, text="Annuler", command=self.cancel).pack(side='right', padx=(5, 0))
        ttk.Button(buttons_frame, text="Sauvegarder", command=self.save_quote).pack(side='right')
    
    def load_clients(self):
        """Load clients into combobox"""
        clients = self.db.get_clients()
        client_names = [f"{client.name} (ID: {client.id})" for client in clients]
        self.client_combo['values'] = client_names
    
    def add_or_update_site(self):
        """Add new site or update existing site based on current mode"""
        if self.is_editing_site:
            self.update_site()
        else:
            self.add_site()
    
    def add_site(self):
        """Add site to quote with detailed address fields"""
        site_number = self.site_number_var.get().strip()
        address = self.site_address_var.get().strip()
        postal_code = self.site_postal_code_var.get().strip()
        city = self.site_city_var.get().strip()
        latitude = self.site_latitude_var.get().strip()
        longitude = self.site_longitude_var.get().strip()
        description = self.site_description_var.get().strip()
        price_str = self.site_price_var.get().strip()
        
        if not site_number:
            messagebox.showerror("Erreur", "Numéro de site obligatoire")
            return
        
        if not description:
            messagebox.showerror("Erreur", "Description obligatoire")
            return
        
        try:
            price = float(price_str.replace(',', '.'))
        except ValueError:
            messagebox.showerror("Erreur", "Prix invalide")
            return
        
        site = SiteItem(
            site_number=site_number,
            address=address,
            postal_code=postal_code,
            city=city,
            latitude=latitude,
            longitude=longitude,
            description=description,
            price_ht=price
        )
        self.quote.sites.append(site)
        
        # Add to tree with formatted address and coordinates
        self.sites_tree.insert('', 'end', values=(
            site.site_number,
            site.full_address,  # Use the formatted full address
            site.coordinates,   # Use the formatted coordinates
            site.description,
            f"{site.price_ht:.2f} €"
        ))
        
        # Clear form after adding
        self.clear_site_form()
        
        # Update totals
        self.update_totals()
    
    def update_site(self):
        """Update existing site with new values"""
        if self.editing_site_index is None or self.editing_site_index >= len(self.quote.sites):
            messagebox.showerror("Erreur", "Erreur lors de la modification du site")
            return
        
        # Validate input fields
        site_number = self.site_number_var.get().strip()
        address = self.site_address_var.get().strip()
        postal_code = self.site_postal_code_var.get().strip()
        city = self.site_city_var.get().strip()
        latitude = self.site_latitude_var.get().strip()
        longitude = self.site_longitude_var.get().strip()
        description = self.site_description_var.get().strip()
        price_str = self.site_price_var.get().strip()
        
        if not site_number:
            messagebox.showerror("Erreur", "Numéro de site obligatoire")
            return
        
        if not description:
            messagebox.showerror("Erreur", "Description obligatoire")
            return
        
        try:
            price = float(price_str.replace(',', '.'))
        except ValueError:
            messagebox.showerror("Erreur", "Prix invalide")
            return
        
        # Update site data in quote
        site = self.quote.sites[self.editing_site_index]
        site.site_number = site_number
        site.address = address
        site.postal_code = postal_code
        site.city = city
        site.latitude = latitude
        site.longitude = longitude
        site.description = description
        site.price_ht = price
        
        # Update tree view - find the corresponding tree item
        children = self.sites_tree.get_children()
        if self.editing_site_index < len(children):
            item_id = children[self.editing_site_index]
            self.sites_tree.item(item_id, values=(
                site.site_number,
                site.full_address,  # Use the formatted full address
                site.coordinates,   # Use the formatted coordinates
                site.description,
                f"{site.price_ht:.2f} €"
            ))
        
        # Exit edit mode
        self.cancel_edit_site()
        
        # Update totals
        self.update_totals()
        
        messagebox.showinfo("Succès", "Site modifié avec succès")
    
    def edit_site(self):
        """Load selected site data into form fields for editing"""
        selection = self.sites_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un site à modifier")
            return
        
        # Get the index of the selected item
        item_id = selection[0]
        children = self.sites_tree.get_children()
        site_index = children.index(item_id)
        
        if site_index >= len(self.quote.sites):
            messagebox.showerror("Erreur", "Site non trouvé")
            return
        
        # Enter edit mode
        self.is_editing_site = True
        self.editing_site_index = site_index
        
        # Load site data into form fields
        site = self.quote.sites[site_index]
        self.site_number_var.set(site.site_number)
        self.site_address_var.set(site.address)
        self.site_postal_code_var.set(site.postal_code)
        self.site_city_var.set(site.city)
        self.site_latitude_var.set(site.latitude)
        self.site_longitude_var.set(site.longitude)
        self.site_description_var.set(site.description)
        self.site_price_var.set(str(site.price_ht))
        
        # Update UI to reflect edit mode
        self.add_site_button.config(text="Mettre à jour Site")
        self.cancel_edit_button.pack(side='right', padx=(0, 5))
    
    def cancel_edit_site(self):
        """Cancel site editing and return to add mode"""
        # Reset edit mode variables
        self.is_editing_site = False
        self.editing_site_index = None
        
        # Clear form fields
        self.clear_site_form()
        
        # Update UI to reflect add mode
        self.add_site_button.config(text="Ajouter Site")
        self.cancel_edit_button.pack_forget()
    
    def clear_site_form(self):
        """Clear all site form fields"""
        self.site_number_var.set("")
        self.site_address_var.set("")
        self.site_postal_code_var.set("")
        self.site_city_var.set("")
        self.site_latitude_var.set("")
        self.site_longitude_var.set("")
        self.site_description_var.set("")
        self.site_price_var.set("")
    
    def remove_site(self):
        """Remove selected site"""
        selection = self.sites_tree.selection()
        if not selection:
            messagebox.showwarning("Attention", "Veuillez sélectionner un site à supprimer")
            return
        
        # Confirm deletion
        if not messagebox.askyesno("Confirmation", "Êtes-vous sûr de vouloir supprimer ce site ?"):
            return
        
        # Get site index
        item_id = selection[0]
        children = self.sites_tree.get_children()
        index = children.index(item_id)
        
        # If we're editing this site, cancel edit mode
        if self.is_editing_site and self.editing_site_index == index:
            self.cancel_edit_site()
        elif self.is_editing_site and self.editing_site_index > index:
            # Adjust edit index if we're deleting a site before the one being edited
            self.editing_site_index -= 1
        
        # Remove from quote
        if 0 <= index < len(self.quote.sites):
            del self.quote.sites[index]
        
        # Remove from tree
        self.sites_tree.delete(item_id)
        
        # Update totals
        self.update_totals()
    
    def add_item(self):
        """Add item to quote (legacy compatibility)"""
        desc = self.item_desc_var.get().strip()
        price_str = self.item_price_var.get().strip()
        
        if not desc:
            messagebox.showerror("Erreur", "Description obligatoire")
            return
        
        try:
            price = float(price_str.replace(',', '.'))
        except ValueError:
            messagebox.showerror("Erreur", "Prix invalide")
            return
        
        item = QuoteItem(description=desc, price_ht=price, quantity=1)
        self.quote.items.append(item)
        
        # Add to tree
        self.items_tree.insert('', 'end', values=(
            item.description,
            f"{item.price_ht:.2f}",
            item.quantity,
            f"{item.total_ht:.2f}"
        ))
        
        # Clear form
        self.item_desc_var.set("")
        self.item_price_var.set("")
        
        # Update totals
        self.update_totals()
    
    def remove_item(self):
        """Remove selected item (legacy compatibility)"""
        selection = self.items_tree.selection()
        if selection:
            # Get item index
            index = self.items_tree.index(selection[0])
            
            # Remove from quote
            if 0 <= index < len(self.quote.items):
                del self.quote.items[index]
            
            # Remove from tree
            self.items_tree.delete(selection[0])
            
            # Update totals
            self.update_totals()
    
    def update_totals(self):
        """Update totals display"""
        total_ht = self.quote.total_ht
        total_tva = self.quote.total_tva
        total_ttc = self.quote.total_ttc
        
        self.totals_label.config(text=f"Total HT: {total_ht:.2f} € | TVA: {total_tva:.2f} € | TTC: {total_ttc:.2f} €")
    
    def save_quote(self):
        """Save quote to database"""
        if not self.client_var.get():
            messagebox.showerror("Erreur", "Veuillez sélectionner un client")
            return
        
        if not self.quote.sites and not self.quote.items:
            messagebox.showerror("Erreur", "Veuillez ajouter au moins un site ou un article")
            return
        
        # Validate and parse quote date
        try:
            day = int(self.quote_day_var.get())
            month = int(self.quote_month_var.get())
            year = int(self.quote_year_var.get())
            quote_date = datetime.date(year, month, day)
            self.quote.quote_date = quote_date
        except (ValueError, TypeError):
            messagebox.showerror("Erreur", "Date invalide. Veuillez vérifier le format JJ/MM/AAAA")
            return
        
        # Set typology
        self.quote.typology = self.typology_var.get()
        
        # Extract client ID from combobox selection
        try:
            client_text = self.client_var.get()
            client_id = int(client_text.split("ID: ")[1].split(")")[0])
            self.quote.client_id = client_id
            
            # Get client info
            self.quote.client = self.db.get_client_by_id(client_id)
            
            # Note: Site information is now handled through the sites list, not individual fields
            # The old site_number and site_address fields are kept for backward compatibility but not used in new quotes
            
            # Generate quote number if new quote
            is_new_quote = self.quote.id is None
            if not self.quote.number:
                self.quote.number = self.db.generate_quote_number(self.quote.client.name)
            
            # Save to database
            quote_id = self.db.save_quote(self.quote)
            
            # Show appropriate success message
            if is_new_quote:
                # This was a new quote
                messagebox.showinfo("Succès", f"Devis créé avec le numéro: {self.quote.number}")
            else:
                # This was an update
                messagebox.showinfo("Succès", f"Devis {self.quote.number} modifié avec succès")
            
            # Call callback to refresh parent
            if self.callback:
                self.callback()
            
            # Close dialog
            self.dialog.destroy()
            
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de la sauvegarde: {str(e)}")
    
    def load_quote_data(self):
        """Load existing quote data into form"""
        # Set client
        if self.quote.client:
            client_text = f"{self.quote.client.name} (ID: {self.quote.client.id})"
            self.client_var.set(client_text)
        
        # Set quote date
        if self.quote.quote_date:
            self.quote_day_var.set(f"{self.quote.quote_date.day:02d}")
            self.quote_month_var.set(f"{self.quote.quote_date.month:02d}")
            self.quote_year_var.set(str(self.quote.quote_date.year))
        
        # Set typology
        if self.quote.typology:
            self.typology_var.set(self.quote.typology)
        
        # Load sites with detailed address information
        for site in self.quote.sites:
            self.sites_tree.insert('', 'end', values=(
                site.site_number,
                site.full_address,  # Use formatted full address
                site.coordinates,   # Use formatted coordinates
                site.description,
                f"{site.price_ht:.2f} €"
            ))
        
        # Load items (legacy compatibility)
        for item in self.quote.items:
            self.items_tree.insert('', 'end', values=(
                item.description,
                f"{item.price_ht:.2f}",
                item.quantity,
                f"{item.total_ht:.2f}"
            ))
        
        # Update totals
        self.update_totals()
    
    def cancel(self):
        """Cancel dialog"""
        self.dialog.destroy()

def main():
    """Main function to run the application"""
    try:
        app = MainApplication()
        app.run()
    except Exception as e:
        print(f"Error starting application: {e}")
        messagebox.showerror("Erreur critique", f"Impossible de démarrer l'application: {e}")

if __name__ == "__main__":
    main()
